#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档列表提取服务：从.doc或.docx文件提取文档项列表
FastAPI服务，提供get_list端口用于Dashboard展示
"""

import os
import re
import logging
import traceback
import tempfile
from datetime import datetime
from typing import Dict, Any, List, Optional
from pathlib import Path

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from pydantic import BaseModel
from docx import Document
import subprocess

# Load environment variables
try:
    from dotenv import load_dotenv
    load_dotenv()
    logger = logging.getLogger(__name__)
    logger.info("✅ 已加载.env环境变量文件")
except ImportError:
    logger = logging.getLogger(__name__)
    logger.warning("⚠️ python-dotenv未安装，继续运行")

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)

# 创建FastAPI应用
app = FastAPI(
    title="文档列表提取服务",
    description="从.doc或.docx文件提取文档项列表，用于Dashboard展示",
    version="1.0.0"
)

class GetListRequest(BaseModel):
    """请求模型 - 文件路径方式"""
    file_path: str

class DocumentItem(BaseModel):
    """文档项模型"""
    id: str
    title: str
    level: int = 1
    type: str = "heading"
    parent_id: Optional[str] = None

class GetListResponse(BaseModel):
    """响应模型"""
    items: List[DocumentItem]
    total_count: int
    success: bool
    message: str
    processing_details: Optional[Dict[str, Any]] = None

class DocumentListExtractor:
    """文档列表提取器"""
    
    def __init__(self):
        self.heading_patterns = [
            r'^([一二三四五六七八九十]+)[、．.]?\s*(.+)$',
            r'^(\d+(?:\.\d+)*)[、．.]?\s*(.+)$',
            r'^[（(]([一二三四五六七八九十]+)[）)]\s*(.+)$',
            r'^([A-Za-z]+)[、．.]?\s*(.+)$',
            r'^[（(](\d+)[）)]\s*(.+)$',
        ]
        logger.info("📋 文档列表提取器初始化完成")
    
    def extract_from_file_path(self, file_path: str) -> List[DocumentItem]:
        """从文件路径提取文档项列表"""
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail=f"文件不存在: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        if file_ext == '.doc':
            docx_path = self._convert_doc_to_docx(file_path)
            return self._extract_from_docx(docx_path)
        elif file_ext == '.docx':
            return self._extract_from_docx(file_path)
        else:
            raise HTTPException(status_code=422, detail=f"不支持的文件格式: {file_ext}")
    
    def extract_from_upload_file(self, upload_file: UploadFile) -> List[DocumentItem]:
        """从上传文件提取文档项列表"""
        try:
            suffix = Path(upload_file.filename or "temp").suffix
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
                temp_file.write(upload_file.file.read())
                temp_path = temp_file.name
            
            try:
                if suffix.lower() == '.doc':
                    docx_path = self._convert_doc_to_docx(temp_path)
                    return self._extract_from_docx(docx_path)
                elif suffix.lower() == '.docx':
                    return self._extract_from_docx(temp_path)
                else:
                    raise HTTPException(status_code=422, detail=f"不支持的文件格式: {suffix}")
            finally:
                os.unlink(temp_path)
                
        except Exception as e:
            logger.error(f"❌ 处理上传文件失败: {e}")
            raise HTTPException(status_code=500, detail=f"处理上传文件失败: {str(e)}")
    
    def _convert_doc_to_docx(self, doc_path: str) -> str:
        """将.doc文件转换为.docx文件"""
        logger.info("🔄 开始DOC到DOCX转换...")
        
        docx_path = doc_path.replace('.doc', '_converted.docx')
        
        try:
            libreoffice_paths = [
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                'libreoffice',
                'soffice',
            ]
            
            libreoffice_cmd = None
            for path in libreoffice_paths:
                try:
                    result = subprocess.run([path, '--version'], 
                                          capture_output=True, 
                                          text=True, 
                                          timeout=10)
                    if result.returncode == 0:
                        libreoffice_cmd = path
                        break
                except (FileNotFoundError, subprocess.TimeoutExpired):
                    continue
            
            if not libreoffice_cmd:
                raise HTTPException(status_code=500, detail="LibreOffice未安装或不可用")
            
            if os.path.exists(docx_path):
                os.remove(docx_path)
            
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'docx',
                '--outdir', os.path.dirname(doc_path),
                doc_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode != 0:
                raise HTTPException(status_code=500, detail=f"LibreOffice转换失败")
            
            expected_docx = doc_path.replace('.doc', '.docx')
            if os.path.exists(expected_docx):
                if expected_docx != docx_path:
                    os.rename(expected_docx, docx_path)
                logger.info(f"✅ 转换成功: {docx_path}")
                return docx_path
            else:
                raise HTTPException(status_code=500, detail="转换后的文件未找到")
                
        except Exception as e:
            logger.error(f"❌ 转换过程中出错: {e}")
            raise HTTPException(status_code=500, detail=f"文件转换失败: {str(e)}")
    
    def _extract_from_docx(self, docx_path: str) -> List[DocumentItem]:
        """从docx文件提取文档项列表"""
        logger.info(f"📄 开始从docx文件提取文档项: {Path(docx_path).name}")
        
        try:
            doc = Document(docx_path)
            items = []
            item_counter = 0
            
            # 提取段落标题
            for para in doc.paragraphs:
                if para.text.strip():
                    item = self._process_paragraph(para, item_counter)
                    if item:
                        items.append(item)
                        item_counter += 1
            
            # 提取表格标题和内容
            for table_idx, table in enumerate(doc.tables):
                table_items = self._process_table(table, item_counter, table_idx)
                items.extend(table_items)
                item_counter += len(table_items)
            
            logger.info(f"✅ 成功提取 {len(items)} 个文档项")
            return items
            
        except Exception as e:
            logger.error(f"❌ 提取文档项失败: {e}")
            raise HTTPException(status_code=500, detail=f"文档解析失败: {str(e)}")
    
    def _process_paragraph(self, para, counter: int) -> Optional[DocumentItem]:
        """处理段落，识别标题和重要内容"""
        text = para.text.strip()
        if not text or len(text) < 2:
            return None
        
        style_name = para.style.name if para.style else ""
        is_heading = False
        level = 1
        
        # 检查标题样式
        if "Heading" in style_name or "标题" in style_name:
            is_heading = True
            level_match = re.search(r'(\d+)', style_name)
            if level_match:
                level = int(level_match.group(1))
        
        # 检查格式（加粗等）
        if para.runs:
            first_run = para.runs[0]
            if first_run.bold:
                is_heading = True
        
        # 通过文本模式识别编号标题
        title_info = self._extract_title_info(text)
        if title_info:
            is_heading = True
            level = title_info['level']
            text = title_info['title']
        
        # 过滤不重要的内容
        if not is_heading and len(text) < 5:
            return None
        
        if self._is_header_footer(text):
            return None
        
        return DocumentItem(
            id=str(counter + 1),
            title=text[:100],
            level=level,
            type="heading" if is_heading else "paragraph"
        )
    
    def _process_table(self, table, start_counter: int, table_idx: int) -> List[DocumentItem]:
        """处理表格，提取表格标题和重要行"""
        items = []
        counter = start_counter
        
        table_title = f"表格 {table_idx + 1}"
        
        if table.rows and table.rows[0].cells:
            first_row_text = " | ".join([cell.text.strip() for cell in table.rows[0].cells if cell.text.strip()])
            if first_row_text and len(first_row_text) > 5:
                table_title = first_row_text[:50] + "..." if len(first_row_text) > 50 else first_row_text
        
        items.append(DocumentItem(
            id=f"{counter + 1}",
            title=table_title,
            level=2,
            type="table"
        ))
        counter += 1
        
        # 提取重要行
        for row_idx, row in enumerate(table.rows[1:], 1):
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            
            if self._is_important_table_row(row_text):
                items.append(DocumentItem(
                    id=f"{counter + 1}",
                    title=row_text[:80] + "..." if len(row_text) > 80 else row_text,
                    level=3,
                    type="table_row",
                    parent_id=f"{start_counter + 1}"
                ))
                counter += 1
        
        return items
    
    def _extract_title_info(self, text: str) -> Optional[Dict[str, Any]]:
        """提取标题信息（编号和级别）"""
        for pattern in self.heading_patterns:
            match = re.match(pattern, text.strip())
            if match:
                groups = match.groups()
                if len(groups) >= 2:
                    number_part = groups[0]
                    title_part = groups[1].strip()
                    level = self._calculate_level(number_part)
                    
                    return {
                        'number': number_part,
                        'title': title_part,
                        'level': level
                    }
        return None
    
    def _calculate_level(self, number_part: str) -> int:
        """根据编号计算层级"""
        if '.' in number_part:
            return len(number_part.split('.'))
        
        chinese_numbers = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        if any(cn in number_part for cn in chinese_numbers):
            return 1
        
        if number_part.isdigit():
            num = int(number_part)
            if num <= 10:
                return 1
            elif num <= 100:
                return 2
            else:
                return 3
        
        return 1
    
    def _is_header_footer(self, text: str) -> bool:
        """判断是否为页眉页脚"""
        patterns = [
            r'第\s*\d+\s*页',
            r'共\s*\d+\s*页',
            r'\d{4}[-/]\d{1,2}[-/]\d{1,2}',
            r'^页\s*\d+',
            r'^\s*\d+\s*$',
        ]
        
        for pattern in patterns:
            if re.search(pattern, text):
                return True
        
        return len(text) < 3 or text.isdigit()
    
    def _is_important_table_row(self, row_text: str) -> bool:
        """判断表格行是否重要"""
        if not row_text or len(row_text.strip()) < 5:
            return False
        
        keywords = [
            '小计', '合计', '总计', '汇总',
            '项目', '工程', '施工', '建设',
            '标准', '规范', '要求', '规定',
            '计划', '方案', '设计', '图纸',
            '质量', '安全', '进度', '费用'
        ]
        
        return any(keyword in row_text for keyword in keywords)

# 全局提取器实例
extractor = DocumentListExtractor()

@app.post("/get_list", response_model=GetListResponse)
async def get_list_endpoint(request: GetListRequest):
    """文档列表提取端点 - 文件路径方式"""
    logger.info("📥 接收到文档列表提取请求（文件路径模式）")
    logger.info(f"   文件路径: {request.file_path}")
    
    try:
        items = extractor.extract_from_file_path(request.file_path)
        
        return GetListResponse(
            items=items,
            total_count=len(items),
            success=True,
            message="文档列表提取成功",
            processing_details={
                "file_path": request.file_path,
                "extraction_time": datetime.now().isoformat(),
                "item_types": {item_type: len([i for i in items if i.type == item_type]) 
                             for item_type in set(item.type for item in items)}
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ 文档列表提取失败: {e}")
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")

@app.post("/get_list_upload", response_model=GetListResponse)
async def get_list_upload_endpoint(file: UploadFile = File(...)):
    """文档列表提取端点 - 文件上传方式（推荐）"""
    logger.info("📥 接收到文档列表提取请求（文件上传模式）")
    logger.info(f"   上传文件: {file.filename}")
    
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="未提供文件名")
        
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in ['.doc', '.docx']:
            raise HTTPException(status_code=422, detail=f"不支持的文件格式: {file_ext}")
        
        items = extractor.extract_from_upload_file(file)
        
        return GetListResponse(
            items=items,
            total_count=len(items),
            success=True,
            message="文档列表提取成功",
            processing_details={
                "original_filename": file.filename,
                "extraction_time": datetime.now().isoformat(),
                "item_types": {item_type: len([i for i in items if i.type == item_type]) 
                             for item_type in set(item.type for item in items)}
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ 文档列表提取失败: {e}")
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")

@app.get("/health")
async def health_check():
    """健康检查端点"""
    return {
        "status": "healthy",
        "service": "文档列表提取服务",
        "timestamp": datetime.now().isoformat()
    }

@app.get("/")
async def root():
    """根端点"""
    return {
        "message": "文档列表提取服务",
        "version": "1.0.0",
        "description": "从.doc或.docx文件提取文档项列表，用于Dashboard展示",
        "features": [
            "支持.doc和.docx文件格式",
            "智能识别标题层级关系",
            "提取表格标题和重要行",
            "保持文档结构和编号",
            "支持文件路径和上传两种方式"
        ],
        "endpoints": {
            "get_list": "POST /get_list - 文档列表提取（文件路径方式）",
            "get_list_upload": "POST /get_list_upload - 文档列表提取（文件上传方式，推荐）",
            "health": "GET /health - 健康检查"
        },
        "supported_formats": [".doc", ".docx"],
        "output_format": "结构化JSON列表，包含id、title、level、type等字段"
    }

if __name__ == "__main__":
    import uvicorn
    
    print("🚀 启动文档列表提取服务...")
    print("📋 服务功能: 从文档中提取项目列表用于Dashboard展示")
    print("🌐 访问地址: http://localhost:8002")
    print("📖 API文档: http://localhost:8002/docs")
    print("📄 支持格式: .doc, .docx")
    print("=" * 50)
    
    uvicorn.run(
        "get_list:app",
        host="0.0.0.0",
        port=8002,
        reload=True,
        log_level="info"
    ) 