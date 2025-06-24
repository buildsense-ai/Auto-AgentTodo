#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI Document Processing MCP Server
Provides tools for document template insertion and document list extraction
"""

import os
import json
import logging
import traceback
import hashlib
import tempfile
import re
import subprocess
from datetime import datetime
from typing import Dict, Any, Optional, Union, List
from pathlib import Path

from fastmcp import FastMCP
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
import fitz  # PyMuPDF
from docx import Document as DocxDocument

# Load environment variables from .env file
try:
    from dotenv import load_dotenv
    load_dotenv()  # 自动加载当前目录下的.env文件
except ImportError:
    pass

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

# Initialize FastMCP server
mcp = FastMCP("AI Document Processing Server 🤖")

# Ensure output directory exists
OUTPUT_DIR = "generated_docs"
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# ==================================
# Common Utilities
# ==================================

def get_api_key() -> str:
    """获取OpenRouter API密钥"""
    api_key = os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        # 检查是否是测试模式
        test_mode = os.environ.get("TEST_MODE", "false").lower() == "true"
        if test_mode:
            logger.warning("⚠️ 测试模式：使用模拟API密钥")
            return "test-api-key-for-testing"
        
        logger.error("❌ 未找到OPENROUTER_API_KEY")
        raise RuntimeError("缺少必需的API密钥配置")
    return api_key

class ProcessingError(Exception):
    """自定义处理异常"""
    def __init__(self, message: str, error_code: str, status_code: int = 500):
        self.message = message
        self.error_code = error_code
        self.status_code = status_code
        super().__init__(self.message)

# ==================================
# Document Template Insertion
# ==================================

class DocumentExtractor:
    """文档内容提取器"""
    
    def __init__(self):
        logger.info("📄 文档提取器初始化完成")
    
    def extract_from_file_path(self, file_path: str) -> str:
        """从文件路径提取内容"""
        if not os.path.exists(file_path):
            raise ProcessingError(
                f"原始文档不存在: {file_path}",
                "FILE_NOT_FOUND",
                404
            )
        return self._extract_content(file_path)
    
    def _extract_content(self, file_path: str) -> str:
        """提取文档内容的核心方法"""
        logger.info(f"📄 开始提取文档内容: {Path(file_path).name}")
        
        content = ""
        
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.docx':
                doc = DocxDocument(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
                
                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            content += f"\n表格行: {row_text}"
            
            elif file_ext == '.pdf':
                doc = fitz.open(file_path)
                for page in doc:
                    content += page.get_text()
                doc.close()
            
            elif file_ext in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            else:
                raise ProcessingError(
                    f"不支持的文件格式: {file_ext}",
                    "UNSUPPORTED_FORMAT",
                    422
                )
            
            if not content.strip():
                raise ProcessingError(
                    "文档内容为空",
                    "EMPTY_DOCUMENT",
                    422
                )
            
            logger.info(f"✅ 成功提取内容，长度: {len(content)} 字符")
            return content.strip()
            
        except ProcessingError:
            raise
        except Exception as e:
            logger.error(f"❌ 提取文档内容失败: {e}")
            raise ProcessingError(
                f"文档内容提取失败: {str(e)}",
                "EXTRACTION_ERROR",
                500
            )

class ContentMerger:
    """内容智能合并器"""
    
    def __init__(self, api_key: str):
        """初始化AI客户端"""
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key,
        )
        self.model = "google/gemini-2.5-pro-preview"
        logger.info("🧠 内容合并器初始化完成")
    
    def merge_content(self, template_json: Dict[str, str], original_content: str) -> Dict[str, str]:
        """使用AI智能合并模板JSON和原始内容"""
        logger.info("🧠 开始AI智能合并...")
        
        # 检查是否是测试模式
        test_mode = os.environ.get("TEST_MODE", "false").lower() == "true"
        if test_mode or self.client.api_key == "test-api-key-for-testing":
            logger.warning("⚠️ 测试模式：使用模拟AI合并")
            return self._mock_merge_content(template_json, original_content)
        
        prompt = f"""
你是一个专业的文档处理AI助手。请根据提供的模板JSON结构和原始文档内容，进行智能合并。

模板JSON结构：
{json.dumps(template_json, ensure_ascii=False, indent=2)}

原始文档内容：
{original_content}

任务要求：
1. 分析模板JSON中每个章节的要求
2. 从原始文档内容中提取相关信息
3. 进行语义匹配和内容整合
4. 生成符合模板结构的内容

输出要求：
- 必须返回JSON格式
- 键名与模板JSON完全一致
- 值为根据原始内容智能生成的具体内容
- 如果原始内容中没有相关信息，请基于合理推测生成内容
- 每个章节内容应该完整、专业、符合实际

请直接返回JSON格式的结果，不要包含任何解释文字。
"""
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
            )
            
            if not response or not response.choices or not response.choices[0].message.content:
                raise ProcessingError(
                    "AI响应无效或为空",
                    "AI_NO_RESPONSE",
                    500
                )
            
            # 提取JSON内容
            response_content = response.choices[0].message.content.strip()
            json_str = self._extract_json_from_response(response_content)
            
            try:
                merged_content = json.loads(json_str)
            except json.JSONDecodeError as e:
                logger.error(f"❌ JSON解析失败: {e}")
                logger.error(f"AI响应内容: {response_content}")
                raise ProcessingError(
                    f"AI返回的内容不是有效的JSON格式: {str(e)}",
                    "AI_INVALID_JSON",
                    422
                )
            
            # 验证合并结果
            if not isinstance(merged_content, dict):
                raise ProcessingError(
                    "AI返回的内容不是字典格式",
                    "AI_INVALID_FORMAT",
                    422
                )
            
            logger.info(f"✅ AI合并成功，生成 {len(merged_content)} 个章节")
            for key, value in merged_content.items():
                preview = str(value)[:100] + "..." if len(str(value)) > 100 else str(value)
                logger.info(f"   📝 {key}: {preview}")
            
            return merged_content
            
        except ProcessingError:
            raise
        except Exception as e:
            logger.error(f"❌ AI合并失败: {e}")
            raise ProcessingError(
                f"AI合并过程中发生错误: {str(e)}",
                "AI_MERGE_ERROR",
                500
            )
    
    def _mock_merge_content(self, template_json: Dict[str, str], original_content: str) -> Dict[str, str]:
        """模拟AI合并（测试模式）"""
        logger.info("🧪 模拟AI合并模式")
        
        merged_content = {}
        content_lines = original_content.split('\n')
        content_preview = ' '.join(content_lines[:5])[:200]
        
        for key, description in template_json.items():
            # 基于原始内容和模板描述生成简单的合并内容
            merged_content[key] = f"""根据原始文档内容生成的{key}章节：

{description}

基于原始文档的相关信息：
{content_preview}

本章节内容已根据模板要求进行智能整合，确保符合工程文档的标准格式和要求。具体内容包括项目的基本情况、技术要求、实施方案等关键信息。

注：此内容由测试模式生成，实际应用中将使用真实AI进行智能合并。"""
        
        logger.info(f"✅ 模拟合并完成，生成 {len(merged_content)} 个章节")
        return merged_content
    
    def _extract_json_from_response(self, response_content: str) -> str:
        """从AI响应中提取JSON内容"""
        # 尝试提取JSON
        if "```json" in response_content:
            start = response_content.find("```json") + 7
            end = response_content.find("```", start)
            if end != -1:
                return response_content[start:end].strip()
            else:
                return response_content[response_content.find("```json") + 7:].strip()
        elif response_content.startswith("{") and response_content.endswith("}"):
            return response_content
        else:
            # 查找JSON对象
            start_idx = response_content.find("{")
            if start_idx != -1:
                brace_count = 0
                for i, char in enumerate(response_content[start_idx:], start_idx):
                    if char == "{":
                        brace_count += 1
                    elif char == "}":
                        brace_count -= 1
                        if brace_count == 0:
                            return response_content[start_idx:i+1]
            return response_content

class DocumentGenerator:
    """文档生成器"""
    
    def __init__(self):
        logger.info("📄 文档生成器初始化完成")
    
    def generate_docx(self, merged_content: Dict[str, str], output_path: str) -> Dict[str, Any]:
        """生成最终的docx文档"""
        logger.info("📄 开始生成docx文档...")
        
        try:
            doc = Document()
            
            # 设置文档标题
            title = doc.add_heading('AI智能合并文档', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加生成时间
            timestamp = datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')
            time_para = doc.add_paragraph(f'生成时间: {timestamp}')
            time_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            doc.add_page_break()
            
            # 添加目录标题
            doc.add_heading('目录', level=1)
            
            # 生成目录
            for i, section_title in enumerate(merged_content.keys(), 1):
                toc_para = doc.add_paragraph(f"{i}. {section_title}")
                toc_para.style = 'List Number'
            
            doc.add_page_break()
            
            # 添加正文内容
            for i, (section_title, section_content) in enumerate(merged_content.items(), 1):
                # 添加章节标题
                heading = doc.add_heading(f"{i}. {section_title}", level=1)
                
                # 添加章节内容
                if isinstance(section_content, str):
                    # 处理多段落内容
                    paragraphs = section_content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = doc.add_paragraph(para_text.strip())
                            para.style = 'Normal'
                elif isinstance(section_content, list):
                    # 处理列表内容
                    for item in section_content:
                        para = doc.add_paragraph(str(item))
                        para.style = 'List Bullet'
                else:
                    # 其他类型转为字符串
                    para = doc.add_paragraph(str(section_content))
                    para.style = 'Normal'
                
                # 添加章节间距
                doc.add_paragraph()
            
            # 添加页脚
            footer_section = doc.sections[0]
            footer = footer_section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "本文档由AI智能合并系统自动生成"
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 保存文档
            doc.save(output_path)
            logger.info(f"✅ 成功生成docx文档: {output_path}")
            
            # 验证文档并返回统计信息
            validation_info = self._validate_docx(output_path)
            
            return {
                "sections_count": len(merged_content),
                "file_size": os.path.getsize(output_path),
                "validation": validation_info
            }
            
        except ProcessingError:
            raise
        except Exception as e:
            logger.error(f"❌ 生成docx文档失败: {e}")
            raise ProcessingError(
                f"文档生成失败: {str(e)}",
                "DOCUMENT_GENERATION_ERROR",
                500
            )
    
    def _validate_docx(self, file_path: str) -> Dict[str, Any]:
        """验证生成的docx文档"""
        try:
            # 尝试打开文档进行验证
            doc = Document(file_path)
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            if paragraph_count == 0:
                raise ProcessingError(
                    "生成的文档为空",
                    "EMPTY_GENERATED_DOCUMENT",
                    500
                )
            
            validation_info = {
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "is_valid": True
            }
            
            logger.info(f"✅ 文档验证通过，包含 {paragraph_count} 个段落")
            return validation_info
            
        except ProcessingError:
            raise
        except Exception as e:
            logger.error(f"❌ 文档验证失败: {e}")
            raise ProcessingError(
                f"生成的文档格式有误: {str(e)}",
                "DOCUMENT_VALIDATION_ERROR",
                500
            )

# ==================================
# Document List Extraction
# ==================================

class DocumentItem:
    """文档项模型"""
    def __init__(self, id: str, title: str, level: int = 1, type: str = "heading", parent_id: Optional[str] = None):
        self.id = id
        self.title = title
        self.level = level
        self.type = type
        self.parent_id = parent_id
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "title": self.title,
            "level": self.level,
            "type": self.type,
            "parent_id": self.parent_id
        }

class DocumentListExtractor:
    """文档列表提取器"""
    
    def __init__(self):
        self.heading_patterns = [
            r'^([一二三四五六七八九十]+)[、．.]?\s*(.+)$',
            r'^(\d+(?:\.\d+)*)[、．.]?\s*(.+)$',
            r'^[（(]([一二三四五六七八九十]+)[）)]\s*(.+)$',
            r'^([A-Za-z]+)[、．.]?\s*(.+)$',
            r'^[（(](\d+)[）)]\s*(.+)$',
        ]
        logger.info("📋 文档列表提取器初始化完成")
    
    def extract_from_file_path(self, file_path: str) -> List[DocumentItem]:
        """从文件路径提取文档项列表"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        if file_ext == '.doc':
            docx_path = self._convert_doc_to_docx(file_path)
            return self._extract_from_docx(docx_path)
        elif file_ext == '.docx':
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
    
    def _convert_doc_to_docx(self, doc_path: str) -> str:
        """将.doc文件转换为.docx文件"""
        logger.info("🔄 开始DOC到DOCX转换...")
        
        docx_path = doc_path.replace('.doc', '_converted.docx')
        
        try:
            libreoffice_paths = [
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                'libreoffice',
                'soffice',
            ]
            
            libreoffice_cmd = None
            for path in libreoffice_paths:
                try:
                    result = subprocess.run([path, '--version'], 
                                          capture_output=True, 
                                          text=True, 
                                          timeout=10)
                    if result.returncode == 0:
                        libreoffice_cmd = path
                        break
                except (FileNotFoundError, subprocess.TimeoutExpired):
                    continue
            
            if not libreoffice_cmd:
                raise RuntimeError("LibreOffice未安装或不可用")
            
            if os.path.exists(docx_path):
                os.remove(docx_path)
            
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'docx',
                '--outdir', os.path.dirname(doc_path),
                doc_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice转换失败")
            
            expected_docx = doc_path.replace('.doc', '.docx')
            if os.path.exists(expected_docx):
                if expected_docx != docx_path:
                    os.rename(expected_docx, docx_path)
                logger.info(f"✅ 转换成功: {docx_path}")
                return docx_path
            else:
                raise RuntimeError("转换后的文件未找到")
                
        except Exception as e:
            logger.error(f"❌ 转换过程中出错: {e}")
            raise RuntimeError(f"文件转换失败: {str(e)}")
    
    def _extract_from_docx(self, docx_path: str) -> List[DocumentItem]:
        """从docx文件提取文档项列表"""
        logger.info(f"📄 开始从docx文件提取文档项: {Path(docx_path).name}")
        
        try:
            doc = Document(docx_path)
            items = []
            item_counter = 0
            
            # 提取段落标题
            for para in doc.paragraphs:
                if para.text.strip():
                    item = self._process_paragraph(para, item_counter)
                    if item:
                        items.append(item)
                        item_counter += 1
            
            # 提取表格标题和内容
            for table_idx, table in enumerate(doc.tables):
                table_items = self._process_table(table, item_counter, table_idx)
                items.extend(table_items)
                item_counter += len(table_items)
            
            logger.info(f"✅ 成功提取 {len(items)} 个文档项")
            return items
            
        except Exception as e:
            logger.error(f"❌ 从docx文件提取失败: {e}")
            raise RuntimeError(f"文档内容提取失败: {str(e)}")
    
    def _process_paragraph(self, para, counter: int) -> Optional[DocumentItem]:
        """处理段落，提取标题信息"""
        text = para.text.strip()
        
        # 过滤页眉页脚等无关内容
        if self._is_header_footer(text):
            return None
        
        # 尝试匹配标题模式
        title_info = self._extract_title_info(text)
        if title_info:
            return DocumentItem(
                id=f"item_{counter}",
                title=title_info['title'],
                level=title_info['level'],
                type="heading"
            )
        
        # 如果是重要段落但不是标题，也包含进来
        if len(text) > 10 and len(text) < 200:
            return DocumentItem(
                id=f"item_{counter}",
                title=text,
                level=3,
                type="paragraph"
            )
        
        return None
    
    def _process_table(self, table, start_counter: int, table_idx: int) -> List[DocumentItem]:
        """处理表格，提取重要行"""
        items = []
        counter = start_counter
        
        for row_idx, row in enumerate(table.rows):
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            
            if self._is_important_table_row(row_text):
                items.append(DocumentItem(
                    id=f"table_{table_idx}_row_{row_idx}",
                    title=row_text,
                    level=2,
                    type="table_row",
                    parent_id=f"table_{table_idx}"
                ))
                counter += 1
        
        if items:
            # 添加表格标题
            table_title = DocumentItem(
                id=f"table_{table_idx}",
                title=f"表格 {table_idx + 1}",
                level=1,
                type="table"
            )
            return [table_title] + items
        
        return []
    
    def _extract_title_info(self, text: str) -> Optional[Dict[str, Any]]:
        """从文本中提取标题信息"""
        for pattern in self.heading_patterns:
            match = re.match(pattern, text)
            if match:
                number_part = match.group(1)
                title_part = match.group(2).strip()
                level = self._calculate_level(number_part)
                
                return {
                    'title': f"{number_part}. {title_part}",
                    'level': level,
                    'number': number_part
                }
        
        return None
    
    def _calculate_level(self, number_part: str) -> int:
        """根据编号计算层级"""
        if re.match(r'^\d+$', number_part):
            return 1
        elif re.match(r'^\d+\.\d+$', number_part):
            return 2
        elif re.match(r'^\d+\.\d+\.\d+$', number_part):
            return 3
        elif number_part in ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']:
            return 1
        elif re.match(r'^[A-Za-z]$', number_part):
            return 2
        else:
            return 2
    
    def _is_header_footer(self, text: str) -> bool:
        """判断是否为页眉页脚"""
        header_footer_patterns = [
            r'第\s*\d+\s*页',
            r'共\s*\d+\s*页',
            r'\d{4}年\d{1,2}月\d{1,2}日',
            r'^页码',
            r'^第.*章$',
        ]
        
        for pattern in header_footer_patterns:
            if re.search(pattern, text):
                return True
        
        return len(text) < 5 or len(text) > 300
    
    def _is_important_table_row(self, row_text: str) -> bool:
        """判断表格行是否重要"""
        if not row_text.strip() or len(row_text) < 5:
            return False
        
        # 过滤明显的表头或分隔行
        if re.match(r'^[\s\-\|]+$', row_text):
            return False
        
        # 包含重要关键词的行
        important_keywords = ['项目', '内容', '要求', '标准', '规范', '方案', '措施']
        for keyword in important_keywords:
            if keyword in row_text:
                return True
        
        return len(row_text) > 10 and len(row_text) < 200

# ==================================
# FastMCP Tools
# ==================================

@mcp.tool()
def insert_template(template_json_input: Union[str, Dict[str, str]], original_file_path: str) -> str:
    """
    AI tool to merge a document with a JSON template to generate a new docx file.
    
    It uses an AI model to intelligently merge content from the original document 
    (e.g., .docx, .pdf, .txt) into the structure defined by the JSON template.

    Args:
        template_json_input: A dictionary or file path for the template JSON.
        original_file_path: The path to the original document file.

    Returns:
        The file path of the generated .docx document.
    """
    logger.info("🚀 Starting template insertion process...")
    
    try:
        # Load template json if a path is provided
        if isinstance(template_json_input, str):
            if not os.path.exists(template_json_input):
                raise FileNotFoundError(f"Template JSON file not found: {template_json_input}")
            with open(template_json_input, 'r', encoding='utf-8') as f:
                template_json = json.load(f)
        else:
            template_json = template_json_input

        # Get API key and initialize components
        api_key = get_api_key()
        extractor = DocumentExtractor()
        merger = ContentMerger(api_key)
        generator = DocumentGenerator()

        # 1. Extract original document content
        original_content = extractor.extract_from_file_path(original_file_path)
        
        # 2. AI intelligent merge
        merged_content = merger.merge_content(template_json, original_content)
        
        # 3. Generate output file path
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"merged_document_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)
        
        # 4. Generate docx document
        generation_info = generator.generate_docx(merged_content, output_path)
        
        logger.info(f"✅ Template insertion process completed successfully. Document saved at: {output_path}")
        return output_path

    except (ProcessingError, FileNotFoundError) as e:
        logger.error(f"❌ Processing failed: {e}")
        raise
    except Exception as e:
        logger.error(f"❌ An unexpected error occurred during template insertion: {e}")
        logger.error(traceback.format_exc())
        raise ProcessingError(f"An unexpected error occurred: {str(e)}", "UNEXPECTED_ERROR", 500)

@mcp.tool()
def extract_document_list(file_path: str) -> List[Dict[str, Any]]:
    """
    AI tool to extract a structured list of items from Word documents (.doc/.docx).
    
    This function processes Word documents and extracts headings, paragraphs, and table
    content to create a structured list suitable for dashboard display or further processing.

    Args:
        file_path: Path to the Word document file (.doc or .docx)

    Returns:
        A list of dictionaries containing document items with id, title, level, type, and parent_id
    """
    logger.info(f"🚀 Starting document list extraction from: {file_path}")
    
    try:
        extractor = DocumentListExtractor()
        items = extractor.extract_from_file_path(file_path)
        
        # Convert DocumentItem objects to dictionaries
        result = [item.to_dict() for item in items]
        
        logger.info(f"✅ Successfully extracted {len(result)} items from document")
        return result
        
    except FileNotFoundError as e:
        logger.error(f"❌ File not found: {e}")
        raise
    except ValueError as e:
        logger.error(f"❌ Invalid file format: {e}")
        raise
    except Exception as e:
        logger.error(f"❌ An unexpected error occurred during extraction: {e}")
        logger.error(traceback.format_exc())
        raise RuntimeError(f"Document extraction failed: {str(e)}")

if __name__ == "__main__":
    print("=" * 70)
    print("🤖 AI Document Processing MCP Server")
    print("=" * 70)
    print("\nAvailable tools:")
    print("1. insert_template - Merge documents with JSON templates")
    print("2. extract_document_list - Extract structured lists from Word documents")
    print("\nStarting MCP server...")
    print("=" * 70)
    
    mcp.run() 