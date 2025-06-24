#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模板插入服务：AI智能合并原始文档与模板JSON
FastAPI服务，提供insert_temp端口
"""

import os
import json
import logging
import traceback
import hashlib
import tempfile
from datetime import datetime
from typing import Dict, Any, Optional, Union
from pathlib import Path

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from pydantic import BaseModel, validator
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
import fitz  # PyMuPDF
from docx import Document as DocxDocument

# 配置日志（优先配置）
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

# Load environment variables from .env file
try:
    from dotenv import load_dotenv
    load_dotenv()  # 自动加载当前目录下的.env文件
    logger.info("✅ 已加载.env环境变量文件")
except ImportError:
    logger.warning("⚠️ python-dotenv未安装，将直接从系统环境变量读取配置")
except Exception as e:
    logger.warning(f"⚠️ 加载.env文件时出现问题: {e}")

def get_api_key() -> str:
    """获取OpenRouter API密钥"""
    api_key = os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        # 检查是否是测试模式
        test_mode = os.environ.get("TEST_MODE", "false").lower() == "true"
        if test_mode:
            logger.warning("⚠️ 测试模式：使用模拟API密钥")
            return "test-api-key-for-testing"
        
        logger.error("❌ 未找到OPENROUTER_API_KEY")
        logger.error("请在.env文件中设置: OPENROUTER_API_KEY=your-api-key-here")
        logger.error("或设置系统环境变量: export OPENROUTER_API_KEY='your-api-key-here'")
        logger.error("或设置TEST_MODE=true进入测试模式")
        raise RuntimeError("缺少必需的API密钥配置")
    return api_key

# 创建FastAPI应用
app = FastAPI(
    title="模板插入服务",
    description="AI智能合并原始文档与模板JSON，生成符合模板的docx文档",
    version="1.0.0"
)

# 确保输出目录存在
OUTPUT_DIR = "generated_docs"
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

class InsertTemplateRequest(BaseModel):
    """请求模型 - 文件路径方式（向后兼容）"""
    template_json: Dict[str, str]
    original_file_path: str
    
    @validator('template_json')
    def validate_template_json(cls, v):
        if not v or not isinstance(v, dict):
            raise ValueError('template_json不能为空且必须是字典格式')
        return v

class InsertTemplateResponse(BaseModel):
    """响应模型"""
    final_doc_path: str
    success: bool
    message: str
    processing_details: Optional[Dict[str, Any]] = None

class ProcessingError(Exception):
    """自定义处理异常"""
    def __init__(self, message: str, error_code: str, status_code: int = 500):
        self.message = message
        self.error_code = error_code
        self.status_code = status_code
        super().__init__(self.message)

class DocumentExtractor:
    """文档内容提取器"""
    
    def __init__(self):
        logger.info("📄 文档提取器初始化完成")
    
    def extract_from_file_path(self, file_path: str) -> str:
        """从文件路径提取内容"""
        if not os.path.exists(file_path):
            raise ProcessingError(
                f"原始文档不存在: {file_path}",
                "FILE_NOT_FOUND",
                404
            )
        return self._extract_content(file_path)
    
    def extract_from_upload_file(self, upload_file: UploadFile) -> str:
        """从上传文件提取内容"""
        try:
            # 创建临时文件
            suffix = Path(upload_file.filename or "temp").suffix
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
                temp_file.write(upload_file.file.read())
                temp_path = temp_file.name
            
            try:
                content = self._extract_content(temp_path)
                return content
            finally:
                # 清理临时文件
                os.unlink(temp_path)
                
        except Exception as e:
            raise ProcessingError(
                f"处理上传文件失败: {str(e)}",
                "UPLOAD_PROCESSING_ERROR",
                422
            )
    
    def _extract_content(self, file_path: str) -> str:
        """提取文档内容的核心方法"""
        logger.info(f"📄 开始提取文档内容: {Path(file_path).name}")
        
        content = ""
        
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.docx':
                doc = DocxDocument(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
                
                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            content += f"\n表格行: {row_text}"
            
            elif file_ext == '.pdf':
                doc = fitz.open(file_path)
                for page in doc:
                    content += page.get_text()
                doc.close()
            
            elif file_ext in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            else:
                raise ProcessingError(
                    f"不支持的文件格式: {file_ext}",
                    "UNSUPPORTED_FORMAT",
                    422
                )
            
            if not content.strip():
                raise ProcessingError(
                    "文档内容为空",
                    "EMPTY_DOCUMENT",
                    422
                )
            
            logger.info(f"✅ 成功提取内容，长度: {len(content)} 字符")
            return content.strip()
            
        except ProcessingError:
            raise
        except Exception as e:
            logger.error(f"❌ 提取文档内容失败: {e}")
            raise ProcessingError(
                f"文档内容提取失败: {str(e)}",
                "EXTRACTION_ERROR",
                500
            )

class ContentMerger:
    """内容智能合并器"""
    
    def __init__(self, api_key: str):
        """初始化AI客户端"""
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key,
        )
        self.model = "google/gemini-2.5-pro-preview"
        logger.info("🧠 内容合并器初始化完成")
    
    def merge_content(self, template_json: Dict[str, str], original_content: str) -> Dict[str, str]:
        """使用AI智能合并模板JSON和原始内容"""
        logger.info("🧠 开始AI智能合并...")
        
        # 检查是否是测试模式
        test_mode = os.environ.get("TEST_MODE", "false").lower() == "true"
        if test_mode or self.client.api_key == "test-api-key-for-testing":
            logger.warning("⚠️ 测试模式：使用模拟AI合并")
            return self._mock_merge_content(template_json, original_content)
        
        prompt = f"""
你是一个专业的文档处理AI助手。请根据提供的模板JSON结构和原始文档内容，进行智能合并。

模板JSON结构：
{json.dumps(template_json, ensure_ascii=False, indent=2)}

原始文档内容：
{original_content}

任务要求：
1. 分析模板JSON中每个章节的要求
2. 从原始文档内容中提取相关信息
3. 进行语义匹配和内容整合
4. 生成符合模板结构的内容

输出要求：
- 必须返回JSON格式
- 键名与模板JSON完全一致
- 值为根据原始内容智能生成的具体内容
- 如果原始内容中没有相关信息，请基于合理推测生成内容
- 每个章节内容应该完整、专业、符合实际

请直接返回JSON格式的结果，不要包含任何解释文字。
"""
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
            )
            
            if not response or not response.choices or not response.choices[0].message.content:
                raise ProcessingError(
                    "AI响应无效或为空",
                    "AI_NO_RESPONSE",
                    500
                )
            
            # 提取JSON内容
            response_content = response.choices[0].message.content.strip()
            json_str = self._extract_json_from_response(response_content)
            
            try:
                merged_content = json.loads(json_str)
            except json.JSONDecodeError as e:
                logger.error(f"❌ JSON解析失败: {e}")
                logger.error(f"AI响应内容: {response_content}")
                raise ProcessingError(
                    f"AI返回的内容不是有效的JSON格式: {str(e)}",
                    "AI_INVALID_JSON",
                    422
                )
            
            # 验证合并结果
            if not isinstance(merged_content, dict):
                raise ProcessingError(
                    "AI返回的内容不是字典格式",
                    "AI_INVALID_FORMAT",
                    422
                )
            
            logger.info(f"✅ AI合并成功，生成 {len(merged_content)} 个章节")
            for key, value in merged_content.items():
                preview = str(value)[:100] + "..." if len(str(value)) > 100 else str(value)
                logger.info(f"   📝 {key}: {preview}")
            
            return merged_content
            
        except ProcessingError:
            raise
        except Exception as e:
            logger.error(f"❌ AI合并失败: {e}")
            raise ProcessingError(
                f"AI合并过程中发生错误: {str(e)}",
                "AI_MERGE_ERROR",
                500
            )
    
    def _mock_merge_content(self, template_json: Dict[str, str], original_content: str) -> Dict[str, str]:
        """模拟AI合并（测试模式）"""
        logger.info("🧪 模拟AI合并模式")
        
        merged_content = {}
        content_lines = original_content.split('\n')
        content_preview = ' '.join(content_lines[:5])[:200]
        
        for key, description in template_json.items():
            # 基于原始内容和模板描述生成简单的合并内容
            merged_content[key] = f"""根据原始文档内容生成的{key}章节：

{description}

基于原始文档的相关信息：
{content_preview}

本章节内容已根据模板要求进行智能整合，确保符合工程文档的标准格式和要求。具体内容包括项目的基本情况、技术要求、实施方案等关键信息。

注：此内容由测试模式生成，实际应用中将使用真实AI进行智能合并。"""
        
        logger.info(f"✅ 模拟合并完成，生成 {len(merged_content)} 个章节")
        return merged_content
    
    def _extract_json_from_response(self, response_content: str) -> str:
        """从AI响应中提取JSON内容"""
        # 尝试提取JSON
        if "```json" in response_content:
            start = response_content.find("```json") + 7
            end = response_content.find("```", start)
            if end != -1:
                return response_content[start:end].strip()
            else:
                return response_content[response_content.find("```json") + 7:].strip()
        elif response_content.startswith("{") and response_content.endswith("}"):
            return response_content
        else:
            # 查找JSON对象
            start_idx = response_content.find("{")
            if start_idx != -1:
                brace_count = 0
                for i, char in enumerate(response_content[start_idx:], start_idx):
                    if char == "{":
                        brace_count += 1
                    elif char == "}":
                        brace_count -= 1
                        if brace_count == 0:
                            return response_content[start_idx:i+1]
                brace_count = 0
                for i, char in enumerate(response_content[start_idx:], start_idx):
                    if char == "{":
                        brace_count += 1
                    elif char == "}":
                        brace_count -= 1
                        if brace_count == 0:
                            return response_content[start_idx:i+1]
            return response_content

class DocumentGenerator:
    """文档生成器"""
    
    def __init__(self):
        logger.info("📄 文档生成器初始化完成")
    
    def generate_docx(self, merged_content: Dict[str, str], output_path: str) -> Dict[str, Any]:
        """生成最终的docx文档"""
        logger.info("📄 开始生成docx文档...")
        
        try:
            doc = Document()
            
            # 设置文档标题
            title = doc.add_heading('AI智能合并文档', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加生成时间
            timestamp = datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')
            time_para = doc.add_paragraph(f'生成时间: {timestamp}')
            time_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            doc.add_page_break()
            
            # 添加目录标题
            doc.add_heading('目录', level=1)
            
            # 生成目录
            for i, section_title in enumerate(merged_content.keys(), 1):
                toc_para = doc.add_paragraph(f"{i}. {section_title}")
                toc_para.style = 'List Number'
            
            doc.add_page_break()
            
            # 添加正文内容
            for i, (section_title, section_content) in enumerate(merged_content.items(), 1):
                # 添加章节标题
                heading = doc.add_heading(f"{i}. {section_title}", level=1)
                
                # 添加章节内容
                if isinstance(section_content, str):
                    # 处理多段落内容
                    paragraphs = section_content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = doc.add_paragraph(para_text.strip())
                            para.style = 'Normal'
                elif isinstance(section_content, list):
                    # 处理列表内容
                    for item in section_content:
                        para = doc.add_paragraph(str(item))
                        para.style = 'List Bullet'
                else:
                    # 其他类型转为字符串
                    para = doc.add_paragraph(str(section_content))
                    para.style = 'Normal'
                
                # 添加章节间距
                doc.add_paragraph()
            
            # 添加页脚
            footer_section = doc.sections[0]
            footer = footer_section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "本文档由AI智能合并系统自动生成"
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 保存文档
            doc.save(output_path)
            logger.info(f"✅ 成功生成docx文档: {output_path}")
            
            # 验证文档并返回统计信息
            validation_info = self._validate_docx(output_path)
            
            return {
                "sections_count": len(merged_content),
                "file_size": os.path.getsize(output_path),
                "validation": validation_info
            }
            
        except ProcessingError:
            raise
        except Exception as e:
            logger.error(f"❌ 生成docx文档失败: {e}")
            raise ProcessingError(
                f"文档生成失败: {str(e)}",
                "DOCUMENT_GENERATION_ERROR",
                500
            )
    
    def _validate_docx(self, file_path: str) -> Dict[str, Any]:
        """验证生成的docx文档"""
        try:
            # 尝试打开文档进行验证
            doc = Document(file_path)
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            if paragraph_count == 0:
                raise ProcessingError(
                    "生成的文档为空",
                    "EMPTY_GENERATED_DOCUMENT",
                    500
                )
            
            validation_info = {
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "is_valid": True
            }
            
            logger.info(f"✅ 文档验证通过，包含 {paragraph_count} 个段落")
            return validation_info
            
        except ProcessingError:
            raise
        except Exception as e:
            logger.error(f"❌ 文档验证失败: {e}")
            raise ProcessingError(
                f"生成的文档格式有误: {str(e)}",
                "DOCUMENT_VALIDATION_ERROR",
                500
            )

class TemplateInserter:
    """模板插入调度器 - 协调各个组件"""
    
    def __init__(self, api_key: str):
        """初始化各个组件"""
        self.extractor = DocumentExtractor()
        self.merger = ContentMerger(api_key)
        self.generator = DocumentGenerator()
        logger.info("🤖 模板插入调度器初始化完成")
    
    def process_from_file_path(self, template_json: Dict[str, str], original_file_path: str) -> Dict[str, Any]:
        """从文件路径处理模板插入（向后兼容）"""
        logger.info("🚀 开始文件路径模式的模板插入处理...")
        
        # 1. 提取原始文档内容
        original_content = self.extractor.extract_from_file_path(original_file_path)
        
        # 2. AI智能合并
        merged_content = self.merger.merge_content(template_json, original_content)
        
        # 3. 生成输出文件路径
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"merged_document_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)
        
        # 4. 生成docx文档
        generation_info = self.generator.generate_docx(merged_content, output_path)
        
        logger.info(f"✅ 模板插入处理完成: {output_path}")
        return {
            "final_doc_path": output_path,
            "generation_info": generation_info,
            "content_summary": {key: len(str(value)) for key, value in merged_content.items()}
        }
    
    def process_from_upload_file(self, template_json: Dict[str, str], upload_file: UploadFile) -> Dict[str, Any]:
        """从上传文件处理模板插入"""
        logger.info(f"🚀 开始上传文件模式的模板插入处理: {upload_file.filename}")
        
        # 1. 提取原始文档内容
        original_content = self.extractor.extract_from_upload_file(upload_file)
        
        # 2. AI智能合并
        merged_content = self.merger.merge_content(template_json, original_content)
        
        # 3. 生成输出文件路径
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        # 使用上传文件名前缀
        file_prefix = Path(upload_file.filename or "upload").stem
        output_filename = f"merged_{file_prefix}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)
        
        # 4. 生成docx文档
        generation_info = self.generator.generate_docx(merged_content, output_path)
        
        logger.info(f"✅ 模板插入处理完成: {output_path}")
        return {
            "final_doc_path": output_path,
            "generation_info": generation_info,
            "content_summary": {key: len(str(value)) for key, value in merged_content.items()},
            "original_filename": upload_file.filename
        }

# 全局处理器实例
template_inserter = None

@app.on_event("startup")
async def startup_event():
    """应用启动时初始化"""
    global template_inserter
    
    try:
        api_key = get_api_key()
        template_inserter = TemplateInserter(api_key)
        logger.info("🚀 模板插入服务启动完成")
    except Exception as e:
        logger.error(f"❌ 服务启动失败: {e}")
        raise

@app.post("/insert_temp", response_model=InsertTemplateResponse)
async def insert_template_endpoint(request: InsertTemplateRequest):
    """
    模板插入端点 - 文件路径方式（向后兼容）
    
    将原始文档内容与模板JSON进行AI智能合并，生成符合模板的docx文档
    """
    logger.info("📥 接收到模板插入请求（文件路径模式）")
    logger.info(f"   模板章节数: {len(request.template_json)}")
    logger.info(f"   原始文档: {request.original_file_path}")
    
    try:
        if template_inserter is None:
            raise HTTPException(status_code=500, detail="服务未正确初始化")
        
        # 处理模板插入
        result = template_inserter.process_from_file_path(
            template_json=request.template_json,
            original_file_path=request.original_file_path
        )
        
        return InsertTemplateResponse(
            final_doc_path=result["final_doc_path"],
            success=True,
            message="模板插入成功完成",
            processing_details=result
        )
        
    except ProcessingError as e:
        logger.error(f"❌ 处理错误: {e.message}")
        raise HTTPException(status_code=e.status_code, detail=e.message)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ 模板插入处理失败: {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"内部处理失败: {str(e)}")

@app.post("/insert_temp_upload", response_model=InsertTemplateResponse)
async def insert_template_upload_endpoint(
    template_json: str = Form(..., description="模板JSON字符串"),
    file: UploadFile = File(..., description="原始文档文件")
):
    """
    模板插入端点 - 文件上传方式（推荐）
    
    支持上传文件进行AI智能合并，更适合主AI接入和分布式部署
    """
    logger.info("📥 接收到模板插入请求（文件上传模式）")
    logger.info(f"   上传文件: {file.filename}")
    logger.info(f"   文件类型: {file.content_type}")
    
    try:
        if template_inserter is None:
            raise HTTPException(status_code=500, detail="服务未正确初始化")
        
        # 解析模板JSON
        try:
            template_json_dict = json.loads(template_json)
        except json.JSONDecodeError as e:
            raise HTTPException(status_code=400, detail=f"模板JSON格式错误: {str(e)}")
        
        # 验证模板JSON
        if not template_json_dict or not isinstance(template_json_dict, dict):
            raise HTTPException(status_code=400, detail="模板JSON不能为空且必须是字典格式")
        
        # 验证文件
        if not file.filename:
            raise HTTPException(status_code=400, detail="未提供文件名")
        
        # 处理模板插入
        result = template_inserter.process_from_upload_file(
            template_json=template_json_dict,
            upload_file=file
        )
        
        return InsertTemplateResponse(
            final_doc_path=result["final_doc_path"],
            success=True,
            message="模板插入成功完成",
            processing_details=result
        )
        
    except ProcessingError as e:
        logger.error(f"❌ 处理错误: {e.message}")
        raise HTTPException(status_code=e.status_code, detail=e.message)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ 模板插入处理失败: {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"内部处理失败: {str(e)}")

@app.get("/download/{filename}")
async def download_file(filename: str):
    """
    文件下载端点
    
    允许下载生成的文档文件
    """
    file_path = os.path.join(OUTPUT_DIR, filename)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.get("/health")
async def health_check():
    """健康检查端点"""
    return {
        "status": "healthy",
        "service": "模板插入服务",
        "timestamp": datetime.now().isoformat()
    }

@app.get("/")
async def root():
    """根端点"""
    return {
        "message": "模板插入服务",
        "version": "2.0.0",
        "description": "AI智能合并原始文档与模板JSON，生成符合模板的docx文档",
        "features": [
            "模块化架构：Extractor + Merger + Generator",
            "多种输入方式：文件路径 + 文件上传",
            "精确异常处理：400/422/404/500",
            "详细处理信息：生成统计和验证结果"
        ],
        "endpoints": {
            "insert_temp": "POST /insert_temp - 模板插入处理（文件路径方式，向后兼容）",
            "insert_temp_upload": "POST /insert_temp_upload - 模板插入处理（文件上传方式，推荐）",
            "download": "GET /download/{filename} - 下载生成的文档",
            "health": "GET /health - 健康检查"
        },
        "supported_formats": {
            "input": [".docx", ".pdf", ".txt", ".md"],
            "output": [".docx"]
        },
        "ai_model": "google/gemini-2.5-pro-preview"
    }

if __name__ == "__main__":
    import uvicorn
    
    # 检查API密钥配置
    try:
        api_key = get_api_key()
        logger.info(f"✅ API密钥配置正确 (长度: {len(api_key)} 字符)")
    except Exception as e:
        print(f"❌ 错误: {e}")
        print("\n配置方法:")
        print("1. 创建.env文件并添加:")
        print("   OPENROUTER_API_KEY=your-api-key-here")
        print("\n2. 或设置系统环境变量:")
        print("   export OPENROUTER_API_KEY='your-api-key-here'")
        exit(1)
    
    print("🚀 启动模板插入服务...")
    print("📋 服务功能: AI智能合并原始文档与模板JSON")
    print("🌐 访问地址: http://localhost:8001")
    print("📖 API文档: http://localhost:8001/docs")
    print("💡 配置方式: 从.env文件或环境变量读取API密钥")
    print("=" * 50)
    
    uvicorn.run(
        "insert_template:app",
        host="0.0.0.0",
        port=8001,
        reload=True,
        log_level="info"
    ) 