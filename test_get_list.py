#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试文档列表提取服务
"""

import requests
import json
import os
from pathlib import Path

# 服务配置
BASE_URL = "http://localhost:8002"

def test_get_list_service():
    """测试文档列表提取服务"""
    print("🧪 开始测试文档列表提取服务")
    print("=" * 50)
    
    # 测试1: 健康检查
    print("\n1️⃣ 测试健康检查...")
    try:
        response = requests.get(f"{BASE_URL}/health")
        if response.status_code == 200:
            print("✅ 健康检查通过")
            print(f"   响应: {response.json()}")
        else:
            print(f"❌ 健康检查失败: {response.status_code}")
    except Exception as e:
        print(f"❌ 健康检查异常: {e}")
    
    # 测试2: 获取根信息
    print("\n2️⃣ 测试根端点...")
    try:
        response = requests.get(f"{BASE_URL}/")
        if response.status_code == 200:
            print("✅ 根端点正常")
            data = response.json()
            print(f"   服务名称: {data.get('message')}")
            print(f"   版本: {data.get('version')}")
            print(f"   支持格式: {data.get('supported_formats')}")
        else:
            print(f"❌ 根端点失败: {response.status_code}")
    except Exception as e:
        print(f"❌ 根端点异常: {e}")
    
    # 测试3: 文件路径方式（需要准备测试文件）
    print("\n3️⃣ 测试文件路径方式...")
    test_file_path = "template_test.doc"
    if os.path.exists(test_file_path):
        try:
            response = requests.post(
                f"{BASE_URL}/get_list",
                json={"file_path": test_file_path}
            )
            if response.status_code == 200:
                result = response.json()
                print("✅ 文件路径方式测试成功")
                print(f"   提取项目数: {result['total_count']}")
                print(f"   处理状态: {result['success']}")
                print("   提取的项目示例:")
                for i, item in enumerate(result['items'][:5]):  # 显示前5个
                    print(f"     {i+1}. [{item['type']}] {item['title'][:50]}...")
            else:
                print(f"❌ 文件路径方式测试失败: {response.status_code}")
                print(f"   错误: {response.text}")
        except Exception as e:
            print(f"❌ 文件路径方式测试异常: {e}")
    else:
        print(f"⚠️ 测试文件不存在: {test_file_path}")
    
    # 测试4: 文件上传方式
    print("\n4️⃣ 测试文件上传方式...")
    if os.path.exists(test_file_path):
        try:
            with open(test_file_path, 'rb') as f:
                files = {'file': (test_file_path, f, 'application/msword')}
                response = requests.post(f"{BASE_URL}/get_list_upload", files=files)
            
            if response.status_code == 200:
                result = response.json()
                print("✅ 文件上传方式测试成功")
                print(f"   提取项目数: {result['total_count']}")
                print(f"   处理状态: {result['success']}")
                print("   提取的项目示例:")
                for i, item in enumerate(result['items'][:5]):
                    print(f"     {i+1}. [{item['type']}] 级别{item['level']}: {item['title'][:50]}...")
                    
                # 显示处理详情
                if result.get('processing_details'):
                    details = result['processing_details']
                    print(f"   原始文件名: {details.get('original_filename')}")
                    print(f"   项目类型统计: {details.get('item_types')}")
            else:
                print(f"❌ 文件上传方式测试失败: {response.status_code}")
                print(f"   错误: {response.text}")
        except Exception as e:
            print(f"❌ 文件上传方式测试异常: {e}")
    else:
        print(f"⚠️ 测试文件不存在: {test_file_path}")
    
    # 测试5: 错误处理
    print("\n5️⃣ 测试错误处理...")
    try:
        # 测试不存在的文件
        response = requests.post(
            f"{BASE_URL}/get_list",
            json={"file_path": "nonexistent_file.docx"}
        )
        if response.status_code == 404:
            print("✅ 不存在文件错误处理正确")
        else:
            print(f"⚠️ 不存在文件错误处理异常: {response.status_code}")
        
        # 测试不支持的文件格式
        response = requests.post(
            f"{BASE_URL}/get_list",
            json={"file_path": "test.txt"}
        )
        if response.status_code == 422:
            print("✅ 不支持格式错误处理正确")
        else:
            print(f"⚠️ 不支持格式错误处理异常: {response.status_code}")
    except Exception as e:
        print(f"❌ 错误处理测试异常: {e}")

def create_sample_test_docx():
    """创建测试用的docx文件"""
    from docx import Document
    
    print("\n📄 创建测试文档...")
    doc = Document()
    
    # 添加标题
    doc.add_heading('工程项目竣工文档清单', 0)
    
    # 添加编号标题
    doc.add_heading('一、施工组织设计', level=1)
    doc.add_paragraph('施工组织设计是指导工程施工的综合性文件。')
    
    doc.add_heading('1.1 施工计划', level=2)
    doc.add_paragraph('详细的施工进度安排和时间计划。')
    
    doc.add_heading('1.2 质量控制', level=2)
    doc.add_paragraph('质量管理体系和控制措施。')
    
    doc.add_heading('二、技术资料', level=1)
    doc.add_paragraph('工程技术资料包括设计图纸、技术规范等。')
    
    # 添加表格
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '文档名称'
    hdr_cells[2].text = '责任人'
    
    # 数据行
    data = [
        ('1', '施工图纸', '张工'),
        ('2', '质量验收标准', '李工'),
        ('3', '安全技术规范', '王工')
    ]
    
    for i, (seq, name, person) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = seq
        row_cells[1].text = name
        row_cells[2].text = person
    
    doc.add_heading('三、竣工验收', level=1)
    doc.add_paragraph('工程竣工验收相关文档。')
    
    doc.add_heading('(一) 验收报告', level=2)
    doc.add_paragraph('工程质量验收报告。')
    
    doc.add_heading('(二) 试验资料', level=2)
    doc.add_paragraph('各项试验检测资料。')
    
    # 保存文档
    test_file = 'test_document.docx'
    doc.save(test_file)
    print(f"✅ 测试文档创建完成: {test_file}")
    return test_file

if __name__ == "__main__":
    print("🚀 文档列表提取服务测试工具")
    print("请确保服务已启动: python get_list.py")
    print("服务地址: http://localhost:8002")
    print()
    
    # 创建测试文档
    test_file = create_sample_test_docx()
    
    # 运行测试
    test_get_list_service()
    
    print("\n" + "="*50)
    print("🎯 测试完成！")
    print("💡 提示:")
    print("  - 如果服务未启动，请先运行: python get_list.py")
    print("  - 如果需要测试.doc文件，请确保安装了LibreOffice")
    print("  - 查看API文档: http://localhost:8002/docs") 