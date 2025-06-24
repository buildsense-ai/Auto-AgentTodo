#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
演示脚本：专门测试用户提到的两种特定情况
1. "项目名称：" (冒号后空白)
2. "致____（监理单位）" (下划线加括号提示)

图片附件功能演示脚本
展示新的图片占位符功能如何工作
"""

import os
import sys
import json
import logging
from datetime import datetime

# 设置路径以便导入主模块
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from main import AIDocGenerator

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

def create_demo_template():
    """创建演示模板，包含所有需要测试的特定情况"""
    from docx import Document
    
    doc = Document()
    
    # 添加标题
    doc.add_heading('特定情况演示模板', 0)
    
    # 添加具体的测试段落
    doc.add_paragraph("项目名称：")  # 情况1：只有冒号，后面空白
    doc.add_paragraph("致____（监理单位）")  # 情况2：下划线加括号
    doc.add_paragraph("致____（施工单位）")  # 情况2的另一个例子，测试空白处理
    doc.add_paragraph("审核人（签字）：")  # 签字字段，应被忽略
    
    # 添加一个简单表格
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "负责人："
    table.cell(0, 1).text = ""
    table.cell(1, 0).text = "项目负责人（签字）：" # 表格中的签字字段
    table.cell(1, 1).text = ""
    
    template_path = "demo_specific_template.docx"
    doc.save(template_path)
    logger.info(f"✅ 创建演示模板: {template_path}")
    return template_path

def create_demo_data():
    """创建演示数据 - 故意不包含某些字段"""
    demo_data = {
        "project_name": "古建筑修缮项目",
        "supervision_company": "中建监理有限公司",
        "project_leader": "王工程师",
        "审核人": "张总" # AI不应填充到签字字段
        # 故意不包含施工单位，测试下划线恢复
    }
    
    data_path = "demo_specific_data.json"
    with open(data_path, 'w', encoding='utf-8') as f:
        json.dump(demo_data, f, ensure_ascii=False, indent=2)
    
    logger.info(f"✅ 创建演示数据: {data_path}")
    return data_path, demo_data

def create_demo_template_with_images():
    """
    创建一个包含图片占位符的演示模板内容
    """
    template_content = """
项目报告

项目名称：古建筑修复工程
项目负责人：张工程师
审核日期：2025-01-20

一、项目概述
本项目旨在对历史建筑进行全面修复。详细的施工图纸请参考：{{image:shiGongTu}}

二、现场情况
经过实地勘察，发现建筑物存在多处损坏。现场照片详见：{{image:xianChangZhaoPian}}

三、损坏评估
建筑物主要损坏部位的详细图像请查看：{{image:sunHuaiTu}}

四、修复方案
基于以上分析，制定了详细的修复计划。设计图纸请参考：{{image:sheJiTu}}

五、结论
本项目将按照既定计划进行修复工作。
    """
    
    print("📄 演示模板内容：")
    print("=" * 50)
    print(template_content)
    print("=" * 50)
    
    return template_content

def create_demo_attachments_map():
    """
    创建演示用的图片附件映射
    """
    attachments_map = {
        "shiGongTu": "uploads/construction_drawing.png",
        "xianChangZhaoPian": "uploads/site_photo.jpg", 
        "sunHuaiTu": "uploads/damage_assessment.png",
        "sheJiTu": "uploads/design_blueprint.pdf"
    }
    
    print("🖼️  演示图片附件映射：")
    print("=" * 50)
    for key, path in attachments_map.items():
        print(f"   {key} -> {path}")
    print("=" * 50)
    
    return attachments_map

def demonstrate_replacement_logic():
    """
    演示图片占位符替换逻辑
    """
    template_content = create_demo_template_with_images()
    attachments_map = create_demo_attachments_map()
    
    # 模拟替换过程
    print("🔄 图片占位符替换过程演示：")
    print("=" * 50)
    
    # 创建引用映射
    attachment_ref_map = {}
    ordered_attachments = list(attachments_map.items())
    for i, (key, _) in enumerate(ordered_attachments):
        attachment_ref_map[key.strip()] = i + 1
    
    print("📋 附件编号映射：")
    for key, number in attachment_ref_map.items():
        print(f"   {key} -> 附件{number}")
    
    print("\n🔄 替换结果：")
    result_content = template_content
    for key, number in attachment_ref_map.items():
        placeholder = f"{{{{image:{key}}}}}"
        replacement = f"（详见附件{number}）"
        result_content = result_content.replace(placeholder, replacement)
        print(f"   '{placeholder}' -> '{replacement}'")
    
    print("\n📄 替换后的文档内容：")
    print("=" * 50)
    print(result_content)
    print("=" * 50)
    
    print("\n📎 文档末尾将添加的附件列表：")
    print("=" * 50)
    print("附件列表")
    print()
    for i, (key, path) in enumerate(ordered_attachments):
        print(f"附件 {i+1}: {key}")
        print(f"   [图片文件: {path}]")
        print()
    print("=" * 50)

def create_usage_instructions():
    """
    创建使用说明
    """
    instructions = """
🎯 图片附件功能使用指南

1. **在Word模板中添加图片占位符**：
   - 使用格式：{{image:描述性键名}}
   - 例如：{{image:shiGongTu}}、{{image:xianChangZhaoPian}}

2. **上传包含图片的文件**：
   - 支持PNG、JPG等图片文件
   - 支持包含图片的PDF文件（系统会自动提取）

3. **AI自动处理**：
   - AI会分析图片内容
   - 为每张图片分配描述性键名
   - 生成attachments_map映射

4. **文档生成结果**：
   - 占位符被替换为"（详见附件N）"
   - 实际图片附加在文档末尾
   - 按顺序编号：附件1、附件2...

5. **优势**：
   - 保持模板文字排版整洁
   - 图片统一管理在文档末尾
   - 自动编号和引用
   - 支持多种图片来源
    """
    
    print(instructions)

def main():
    """
    主演示函数 - 图片附件功能
    """
    print("🚀 AI文档生成器 - 图片附件功能演示")
    print("=" * 60)
    
    # 演示替换逻辑
    demonstrate_replacement_logic()
    
    print("\n" + "=" * 60)
    
    # 显示使用说明
    create_usage_instructions()
    
    print("\n✅ 演示完成！")
    print("💡 提示：现在您可以在Word模板中使用 {{image:键名}} 占位符了！")

def main_legacy():
    """演示主函数 - 原有功能"""
    logger.info("🎯 开始特定情况演示")
    logger.info("=" * 60)
    logger.info("演示内容：")
    logger.info("1. '项目名称：' -> '项目名称：古建筑修缮项目'")
    logger.info("2. '致____（监理单位）' -> '致中建监理有限公司（监理单位）'")
    logger.info("3. '致____（施工单位）' -> '致____（施工单位）' (恢复下划线)")
    logger.info("4. '审核人（签字）：' -> 保持原样，不填充")
    logger.info("=" * 60)
    
    try:
        # 获取API密钥
        api_key = os.getenv('OPENROUTER_API_KEY')
        if not api_key:
            logger.error("❌ 未找到 OPENROUTER_API_KEY 环境变量")
            return False
        
        # 创建演示文件
        template_path = create_demo_template()
        data_path, demo_data = create_demo_data()
        
        # 初始化AI生成器
        from main import AIDocGenerator
        generator = AIDocGenerator(api_key)
        
        # 创建输出路径
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = f"demo_specific_output_{timestamp}.docx"
        
        # 运行生成过程
        logger.info("🚀 开始演示生成...")
        success = generator.run_generation(
            doc_template_path=template_path,
            output_path=output_path,
            direct_json_data=demo_data
        )
        
        if success:
            logger.info(f"✅ 演示完成！输出文件: {output_path}")
            logger.info("📋 请检查输出文件验证以下内容：")
            logger.info("   ✓ '项目名称：' 后应该有具体项目名称")
            logger.info("   ✓ '致____（监理单位）' 应该填入具体监理单位")
            logger.info("   ✓ '致____（施工单位）' 应该恢复为下划线")
            logger.info("   ✓ '审核人（签字）：' 和表格中的签字字段应保持原样，不被填充")
            logger.info("   ✓ 表格中的'负责人'应正确填写")
            
            # 清理演示文件
            if os.path.exists(template_path):
                os.remove(template_path)
            if os.path.exists(data_path):
                os.remove(data_path)
            logger.info("🧹 清理临时演示文件")
            
            return True
        else:
            logger.error("❌ 演示失败")
            return False
            
    except Exception as e:
        logger.error(f"❌ 演示过程中发生错误: {e}", exc_info=True)
        return False

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == "--legacy":
        success = main_legacy()
    sys.exit(0 if success else 1) 
    else:
        main() 